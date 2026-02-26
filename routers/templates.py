# marble_app/routers/templates.py
from __future__ import annotations

import os
import json
import hashlib
import mimetypes
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Literal, Optional, Tuple

from fastapi import (
    APIRouter,
    UploadFile,
    File,
    Form,
    HTTPException,
    Body,
    Request,
    Response,
)
from fastapi.responses import JSONResponse, Response as FastAPIResponse
from pydantic import BaseModel, Field

from docx import Document
import mammoth
import requests

from .storage_s3 import (
    tenant_data_prefix,
    list_keys,
    get_bytes,
    put_bytes,
    delete_key,
)

# -------------------- env --------------------
DOCSERVICE_URL = os.environ.get("ONLYOFFICE_DOCSERVICE_URL", "").strip() or "http://localhost:8082"
PUBLIC_BASE_URL = os.environ.get("PUBLIC_BASE_URL", "").strip() or "http://127.0.0.1:8000"

# -------------------- Routers --------------------
router = APIRouter(prefix="/contracts", tags=["contracts"])
onlyoffice_router = APIRouter(prefix="/onlyoffice", tags=["onlyoffice"])

# -------------------- S3 layout --------------------
# templates live under: tenants/<slug>/data/Templates/
TEMPLATES_PREFIX = tenant_data_prefix("data", "Templates").rstrip("/") + "/"
INDEX_KEY = TEMPLATES_PREFIX + "templates_index.json"
VERSIONS_PREFIX = TEMPLATES_PREFIX + "Versions/"  # optional backups
MAPPINGS_PREFIX = TEMPLATES_PREFIX  # store {tid}.mapping.json alongside templates

# -------------------- Models --------------------
TypeStr = Literal["author", "illustrator", "generic"]
DealTypeStr = Literal["single", "series", "other"]

class ContractTemplate(BaseModel):
    id: str
    name: str
    filename: str
    type: TypeStr
    dealType: DealTypeStr
    uploadedAt: str
    placeholders: Optional[List[str]] = None
    previewUrl: Optional[str] = None

class MappingPosition(BaseModel):
    x_pct: float = Field(ge=0, le=100)
    y_pct: float = Field(ge=0, le=100)

class MappingItem(BaseModel):
    field_key: str
    placeholder: str
    position: MappingPosition

class TemplateMapping(BaseModel):
    templateId: str
    mapping: List[MappingItem]
    _format: Optional[str] = None

class InsertTokenPayload(BaseModel):
    placeholder: str
    paragraph_index: int
    char_offset: int
    left_ctx: Optional[str] = None
    right_ctx: Optional[str] = None

# -------------------- Helpers --------------------
def _abs_url(request: Request, rel: str) -> str:
    base = (PUBLIC_BASE_URL or str(request.base_url)).rstrip("/")
    if not rel.startswith("/"):
        rel = "/" + rel
    return f"{base}{rel}"

def _read_index() -> List[ContractTemplate]:
    try:
        raw = get_bytes(INDEX_KEY)
    except Exception:
        return []
    try:
        data = json.loads(raw.decode("utf-8"))
        if isinstance(data, list):
            return [ContractTemplate(**item) for item in data]
        return []
    except Exception:
        return []

def _write_index(items: List[ContractTemplate]) -> None:
    blob = json.dumps([i.model_dump() for i in items], indent=2, ensure_ascii=False).encode("utf-8")
    put_bytes(INDEX_KEY, blob, content_type="application/json")

def _get_template_obj(tid: str) -> Optional[ContractTemplate]:
    for it in _read_index():
        if it.id == tid:
            return it
    for it in _synthesize_templates_from_s3():
        if it.id == tid:
            return it
    return None


def _synthesize_templates_from_s3() -> List[ContractTemplate]:
    """When templates_index.json is missing or empty, list .docx keys in S3 and build template list."""
    out: List[ContractTemplate] = []
    try:
        keys = list_keys(TEMPLATES_PREFIX)
    except Exception:
        return out
    index_basename = "templates_index.json"
    for key in keys:
        if not key or not key.lower().endswith(".docx"):
            continue
        # Skip files under Versions/ or anything that is the index
        parts = key.split("/")
        basename = parts[-1] if parts else key
        if basename == index_basename or "Versions/" in key:
            continue
        # id: prefix before first _ (e.g. abc123_My_Template.docx -> abc123); else full stem
        stem = basename[:-5] if basename.lower().endswith(".docx") else basename
        if "_" in stem:
            tid, name_part = stem.split("_", 1)
            name = name_part.replace("_", " ").strip() or tid
        else:
            tid = stem
            name = stem
        out.append(
            ContractTemplate(
                id=tid,
                name=name,
                filename=basename,
                type="generic",
                dealType="other",
                uploadedAt=datetime.utcnow().isoformat() + "Z",
                placeholders=None,
                previewUrl=None,
            )
        )
    return out

def _template_key_from_obj(it: ContractTemplate) -> str:
    return TEMPLATES_PREFIX + it.filename

def _key_for_id(tid: str) -> str:
    it = _get_template_obj(tid)
    if it and it.filename:
        return _template_key_from_obj(it)

    # fallback scan (slower)
    keys = list_keys(TEMPLATES_PREFIX)
    for k in keys:
        name = k.split("/")[-1]
        if name.startswith(tid + "_") or name.startswith(tid + "__"):
            return k
    return TEMPLATES_PREFIX + f"{tid}_MISSING"

def _mapping_key(tid: str) -> str:
    return MAPPINGS_PREFIX + f"{tid}.mapping.json"

def _version_key(tid: str, ts: str) -> str:
    return VERSIONS_PREFIX + f"{tid}/{ts}.docx"

def _touch_index_after_save(tid: str, filename: str, placeholders: Optional[List[str]] = None) -> None:
    items = _read_index()
    for i, it in enumerate(items):
        if it.id == tid:
            it.uploadedAt = datetime.utcnow().isoformat() + "Z"
            it.filename = filename
            if placeholders is not None:
                it.placeholders = placeholders
            items[i] = it
            _write_index(items)
            return

def _download_docx_to_tmp(key: str) -> str:
    data = get_bytes(key)
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    with open(path, "wb") as f:
        f.write(data)
    return path

def _upload_docx_from_path(key: str, path: str) -> None:
    with open(path, "rb") as f:
        data = f.read()
    put_bytes(
        key,
        data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ---- Insert token helpers ----
def _clone_run_format(src_run, dst_run):
    sf, df = src_run.font, dst_run.font
    df.name = sf.name
    df.size = sf.size
    df.bold = sf.bold
    df.italic = sf.italic
    df.underline = sf.underline
    df.color.rgb = getattr(sf.color, "rgb", None)
    df.all_caps = sf.all_caps
    df.small_caps = sf.small_caps

def _normalize_text(s: str) -> str:
    return (s or "").replace("\u00A0", " ").replace("\r", "").replace("\n", " ").replace("\t", " ").strip()

def _find_offset_with_context(full_text: str, raw_offset: int, left_ctx: Optional[str], right_ctx: Optional[str]) -> int:
    ntext = _normalize_text(full_text)
    left_ctx = _normalize_text(left_ctx) if left_ctx else None
    right_ctx = _normalize_text(right_ctx) if right_ctx else None

    if left_ctx and right_ctx:
        li = ntext.find(left_ctx)
        if li >= 0:
            start = li + len(left_ctx)
            if right_ctx == "" or ntext.find(right_ctx, start) >= 0:
                return min(max(start, 0), len(ntext))
    if left_ctx and not right_ctx:
        li = ntext.find(left_ctx)
        if li >= 0:
            return min(max(li + len(left_ctx), 0), len(ntext))
    if right_ctx and not left_ctx:
        ri = ntext.find(right_ctx)
        if ri >= 0:
            return min(max(ri, 0), len(ntext))
    return min(max(raw_offset, 0), len(ntext))

def _insert_text_at_paragraph_offset(p, offset: int, token: str) -> Tuple[bool, int]:
    runs = list(p.runs)
    positions = []
    total = 0
    for r in runs:
        t = r.text or ""
        start = total
        end = start + len(t)
        positions.append((r, t, start, end))
        total = end

    offset = max(0, min(offset, total))

    if not runs:
        p.add_run(token)
        return True, 0

    target_index = None
    for i, (_, _, start, end) in enumerate(positions):
        if start <= offset <= end:
            target_index = i
            break

    if target_index is None:
        r_last, t_last, *_ = positions[-1]
        r_last.text = (t_last or "") + token
        return True, total

    r, t, start, end = positions[target_index]
    inner = max(0, min(offset - start, len(t)))
    before = t[:inner]
    after = t[inner:]

    r.text = before
    new_run = p.add_run("")
    p._p.remove(new_run._r)
    r._r.addnext(new_run._r)
    new_run.text = token
    _clone_run_format(r, new_run)

    if after:
        tail_run = p.add_run("")
        p._p.remove(tail_run._r)
        new_run._r.addnext(tail_run._r)
        tail_run.text = after
        _clone_run_format(r, tail_run)

    return True, offset

# -------------------- Routes --------------------
@router.get("/templates", response_model=List[ContractTemplate])
async def list_templates(request: Request):
    items = _read_index()
    from_index = bool(items)
    if not items:
        items = _synthesize_templates_from_s3()
    base = str(request.base_url).rstrip("/")
    changed = False
    for it in items:
        want = f"{base}/api/contracts/templates/{it.id}/file"
        if it.previewUrl != want:
            it.previewUrl = want
            changed = True
    if changed and from_index:
        _write_index(items)
    return items

@router.post("/templates", response_model=ContractTemplate)
async def upload_template(
    request: Request,
    file: UploadFile = File(...),
    name: str = Form(...),
    type: TypeStr = Form(...),
    dealType: DealTypeStr = Form(...),
):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are accepted.")

    tid = os.urandom(16).hex()
    original_name = (file.filename or "template.docx").rsplit("\\", 1)[-1].rsplit("/", 1)[-1]
    stored_filename = f"{tid}_{original_name}"
    key = TEMPLATES_PREFIX + stored_filename

    data = await file.read()
    put_bytes(
        key,
        data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    item = ContractTemplate(
        id=tid,
        name=name,
        filename=stored_filename,
        type=type,
        dealType=dealType,
        uploadedAt=datetime.utcnow().isoformat() + "Z",
        placeholders=None,
        previewUrl=f"{str(request.base_url).rstrip('/')}/api/contracts/templates/{tid}/file",
    )

    items = _read_index()
    items.append(item)
    _write_index(items)
    return item

@router.get("/templates/{tid}", response_model=ContractTemplate)
async def get_template(tid: str, request: Request):
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")

    want = f"{str(request.base_url).rstrip('/')}/api/contracts/templates/{tid}/file"
    if it.previewUrl != want:
        items = _read_index()
        for i, x in enumerate(items):
            if x.id == tid:
                items[i].previewUrl = want
                break
        _write_index(items)
        it.previewUrl = want
    return it

@router.get("/templates/{tid}/file")
async def get_template_file(tid: str):
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")

    key = _template_key_from_obj(it)
    try:
        data = get_bytes(key)
    except Exception:
        raise HTTPException(status_code=404, detail="File not found in S3")

    media_type = mimetypes.guess_type(it.filename)[0] or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    headers = {
        "Content-Disposition": f'inline; filename="{it.filename.split("_",1)[-1]}"',
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0",
        "ETag": hashlib.sha256(data).hexdigest()[:16],
    }
    return FastAPIResponse(content=data, media_type=media_type, headers=headers)

@router.get("/templates/{tid}/html")
async def get_template_html(tid: str):
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")

    key = _template_key_from_obj(it)
    tmp_path = _download_docx_to_tmp(key)
    try:
        with open(tmp_path, "rb") as f:
            result = mammoth.convert_to_html(f, style_map="")
        raw_html = result.value or ""
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

    parts = raw_html.split("<p")
    rebuilt = []
    p_index = -1
    for i, part in enumerate(parts):
        if i == 0:
            rebuilt.append(part)
            continue
        p_index += 1
        rebuilt.append(f'<p data-p-index="{p_index}"' + part)
    indexed_html = "".join(rebuilt)

    return {"ok": True, "html": indexed_html, "updated": datetime.utcnow().isoformat() + "Z"}

@router.post("/templates/{tid}/insert-token")
async def insert_token(tid: str, body: InsertTokenPayload):
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")

    token_text = (body.placeholder or "").strip()
    if not token_text:
        raise HTTPException(status_code=400, detail="placeholder cannot be empty")

    key = _template_key_from_obj(it)
    tmp_path = _download_docx_to_tmp(key)

    try:
        doc = Document(tmp_path)
        paragraphs = list(doc.paragraphs)
        if body.paragraph_index < 0 or body.paragraph_index >= len(paragraphs):
            raise HTTPException(status_code=400, detail="paragraph_index out of range")

        p = paragraphs[body.paragraph_index]
        para_text = p.text or ""
        effective_offset = _find_offset_with_context(para_text, body.char_offset, body.left_ctx, body.right_ctx)
        ok, eff = _insert_text_at_paragraph_offset(p, effective_offset, token_text)
        if not ok:
            raise HTTPException(status_code=500, detail="failed to insert token")

        doc.save(tmp_path)
        _upload_docx_from_path(key, tmp_path)
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

    # update placeholders in index
    items = _read_index()
    phs = None
    for x in items:
        if x.id == tid:
            cur = set(x.placeholders or [])
            if token_text not in cur:
                cur.add(token_text)
            phs = sorted(cur)
            x.placeholders = phs
            x.uploadedAt = datetime.utcnow().isoformat() + "Z"
            break
    _write_index(items)

    return {
        "ok": True,
        "templateId": tid,
        "paragraph_index": body.paragraph_index,
        "char_offset": eff,
        "inserted": token_text,
        "matched_by": "context" if (body.left_ctx or body.right_ctx) else "offset",
        "saved_at": datetime.utcnow().isoformat() + "Z",
    }

@router.post("/templates/{tid}/mapping")
async def save_template_mapping(tid: str, payload: TemplateMapping):
    if payload.templateId != tid:
        raise HTTPException(status_code=400, detail="templateId mismatch")

    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")

    mapping_obj: Dict[str, Any] = {
        "template_id": tid,
        "saved_at": datetime.utcnow().isoformat() + "Z",
        "mapping": [mi.model_dump() for mi in payload.mapping],
        "_format": payload._format or "docx-overlay",
    }
    put_bytes(_mapping_key(tid), json.dumps(mapping_obj, indent=2, ensure_ascii=False).encode("utf-8"), content_type="application/json")

    ph = sorted({m.placeholder for m in payload.mapping if m.placeholder})
    _touch_index_after_save(tid, it.filename, placeholders=ph)
    return {"ok": True, "templateId": tid, "placeholders": ph}

@router.post("/templates/{tid}/save")
async def persist_updated_template(tid: str, request: Request):
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")

    body = await request.body()
    if not body:
        raise HTTPException(status_code=400, detail="Empty body")

    key = _template_key_from_obj(it)

    # optional backup
    try:
        old = get_bytes(key)
        ts = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
        put_bytes(_version_key(tid, ts), old, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception:
        pass

    put_bytes(key, body, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    _touch_index_after_save(tid, it.filename)
    return {"ok": True, "id": tid, "key": key}

@router.post("/templates/{tid}/force-save")
async def force_save_noop(tid: str):
    return {"ok": True, "id": tid}

@router.delete("/templates/{tid}", status_code=204)
async def delete_template(tid: str):
    items = _read_index()
    idx = next((i for i, it in enumerate(items) if it.id == tid), None)
    if idx is None:
        raise HTTPException(status_code=404, detail="Template not found")

    it = items[idx]
    # delete docx + mapping (versions left as historical unless you want to delete prefix)
    try:
        delete_key(_template_key_from_obj(it))
    except Exception:
        pass
    try:
        delete_key(_mapping_key(tid))
    except Exception:
        pass

    del items[idx]
    _write_index(items)
    return Response(status_code=204)

# -------------------- ONLYOFFICE --------------------
def _doc_key_for_bytes(b: bytes, filename: str) -> str:
    raw = hashlib.sha256(b + filename.encode("utf-8")).hexdigest()
    return raw[:40]

@onlyoffice_router.get("/config/{tid}")
async def onlyoffice_config(tid: str, request: Request):
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")

    key = _template_key_from_obj(it)
    data = get_bytes(key)
    doc_url = _abs_url(request, f"/api/contracts/templates/{tid}/file")
    callback_url = _abs_url(request, f"/api/onlyoffice/callback/{tid}")
    title = it.filename.split("_", 1)[-1]
    doc_key = _doc_key_for_bytes(data, it.filename)

    config: Dict[str, Any] = {
        "document": {
            "fileType": "docx",
            "key": doc_key,
            "title": title,
            "url": doc_url,
            "permissions": {
                "edit": True,
                "download": True,
                "print": True,
                "copy": True,
                "fillForms": True,
                "review": True,
                "comment": True,
            },
        },
        "editorConfig": {
            "mode": "edit",
            "callbackUrl": callback_url,
            "user": {"id": "user", "name": "Template Editor"},
            "autosave": True,
            "customization": {
                "toolbar": True,
                "forcesave": True,
                "hideRightMenu": True,
                "toolbarNoTabs": False,
                "chat": False,
            },
        },
    }
    return {"docServiceUrl": DOCSERVICE_URL.rstrip("/"), "config": config}

@onlyoffice_router.post("/callback/{tid}")
async def onlyoffice_callback(tid: str, payload: Dict[str, Any] = Body(default={})):
    it = _get_template_obj(tid)
    if not it:
        return {"error": 0}

    status = int(payload.get("status") or 0)
    if status not in (2, 6):
        return {"error": 0}

    download_url = payload.get("url")
    if not download_url:
        return {"error": 1}

    headers: Dict[str, str] = {}
    token = payload.get("token")
    if isinstance(token, str) and token:
        headers["Authorization"] = f"Bearer {token}"

    try:
        r = requests.get(download_url, headers=headers, timeout=60)
        r.raise_for_status()
        put_bytes(
            _template_key_from_obj(it),
            r.content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        _touch_index_after_save(tid, it.filename)
    except Exception:
        return {"error": 1}

    return {"error": 0}

# -------------------- WOPI (Collabora) for templates --------------------
try:
    from app.core.config import get_settings
    from app.wopi.tokens import make_wopi_token, verify_wopi_token
except Exception:
    get_settings = None
    make_wopi_token = None
    verify_wopi_token = None


def _wopi_settings():
    if get_settings is None:
        raise HTTPException(status_code=500, detail="WOPI not configured")
    return get_settings()


def _extract_wopi_token(request: Request) -> str:
    token = request.query_params.get("access_token")
    if token:
        return token
    auth = request.headers.get("Authorization")
    if auth and auth.startswith("Bearer "):
        return auth[7:].strip()
    return request.headers.get("X-WOPI-AccessToken") or ""


wopi_templates_router = APIRouter(prefix="/wopi/templates", tags=["WOPI Templates"])


@wopi_templates_router.get("/files/{tid}")
def wopi_template_check_file_info(tid: str, request: Request):
    token = _extract_wopi_token(request)
    settings = _wopi_settings()
    payload = verify_wopi_token(token, settings.wopi_access_token_secret)
    if str(payload.get("file_id")) != str(tid):
        raise HTTPException(status_code=401, detail="Invalid WOPI token")
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="File not found")
    key = _template_key_from_obj(it)
    try:
        data = get_bytes(key)
        size = len(data)
    except Exception:
        raise HTTPException(status_code=404, detail="File not found")
    base_name = it.filename or "template.docx"
    user_id = (payload.get("user_id") or "inksuite-user").strip()
    user_id = "".join(c if c.isalnum() else "_" for c in user_id)[:64] or "inksuite_user"
    return JSONResponse({
        "BaseFileName": base_name,
        "Size": size,
        "OwnerId": "inksuite",
        "UserId": user_id,
        "UserFriendlyName": payload.get("user_id") or "InkSuite User",
        "Version": it.uploadedAt or str(id(it)),
        "SupportsUpdate": True,
        "SupportsLocks": True,
        "SupportsGetLock": True,
        "SupportsPutFile": True,
        "SupportsRename": False,
        "ReadOnly": False,
        "UserCanWrite": True,
    })


@wopi_templates_router.get("/files/{tid}/contents")
def wopi_template_get_file(tid: str, request: Request):
    token = _extract_wopi_token(request)
    settings = _wopi_settings()
    payload = verify_wopi_token(token, settings.wopi_access_token_secret)
    if str(payload.get("file_id")) != str(tid):
        raise HTTPException(status_code=401, detail="Invalid WOPI token")
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")
    key = _template_key_from_obj(it)
    try:
        data = get_bytes(key)
    except Exception:
        raise HTTPException(status_code=404, detail="File not found")
    return FastAPIResponse(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'inline; filename="{it.filename}"'},
    )


@wopi_templates_router.post("/files/{tid}/contents")
async def wopi_template_put_file(tid: str, request: Request):
    token = _extract_wopi_token(request)
    settings = _wopi_settings()
    payload = verify_wopi_token(token, settings.wopi_access_token_secret)
    if str(payload.get("file_id")) != str(tid):
        raise HTTPException(status_code=401, detail="Invalid WOPI token")
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")
    body = await request.body()
    key = _template_key_from_obj(it)
    put_bytes(key, body, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    _touch_index_after_save(tid, it.filename)
    return JSONResponse({"LastModifiedTime": datetime.utcnow().isoformat() + "Z"})


@router.get("/templates/{tid}/collabora-config")
def collabora_config_for_template(tid: str):
    it = _get_template_obj(tid)
    if not it:
        raise HTTPException(status_code=404, detail="Template not found")
    settings = _wopi_settings()
    token = make_wopi_token(
        file_id=tid,
        user_id="inksuite-user",
        ttl=settings.wopi_token_ttl_sec,
        secret=settings.wopi_access_token_secret,
    )
    from urllib.parse import quote
    wopi_src = f"{settings.public_base_url}/api/wopi/templates/files/{quote(tid, safe='')}"
    base = settings.collabora_code_url
    if "/loleaflet" in base or "/browser" in base or base.endswith(".html"):
        editor_url = base
    else:
        editor_url = f"{base}/loleaflet/dist/loleaflet.html"
    return JSONResponse({
        "editorUrl": editor_url,
        "wopiSrc": wopi_src,
        "accessToken": token,
        "readOnly": False,
    })


# aliases
@router.get("/templates/{tid}/onlyoffice-config")
async def alias_onlyoffice_config(tid: str, request: Request):
    return await onlyoffice_config(tid, request)

@router.post("/templates/{tid}/onlyoffice-callback")
async def alias_onlyoffice_callback(tid: str, payload: Dict[str, Any] = Body(default={})):
    return await onlyoffice_callback(tid, payload)