# marble_app/routers/contract_docs.py
from __future__ import annotations

import json
import os
import re
import time
import uuid
import tempfile
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import quote

import boto3
from botocore.config import Config
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, Response as FastAPIResponse
from pydantic import BaseModel, EmailStr

from app.core.db import db_conn

try:
    from docx import Document
    from docx.shared import RGBColor
except Exception as e:
    raise RuntimeError("python-docx not installed. pip install python-docx") from e

try:
    from .storage_s3 import tenant_data_prefix, list_keys, get_bytes, put_bytes

    _S3_AVAILABLE = True
except Exception:
    _S3_AVAILABLE = False

try:
    from app.core.config import get_settings
    from app.wopi.tokens import make_wopi_token, verify_wopi_token
except Exception:
    get_settings = None
    make_wopi_token = None
    verify_wopi_token = None


class AgentInviteIn(BaseModel):
    name: str = ""
    email: EmailStr


class GenerateRequest(BaseModel):
    dealMemo: Dict[str, Any]
    templateId: str
    mapping: Optional[Dict[str, str]] = None


router = APIRouter(prefix="/contracts", tags=["contracts"])
wopi_router = APIRouter(prefix="/wopi", tags=["WOPI"])

# ----------------------- S3 -----------------------
_DRAFTS_AWS_REGION = (os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "us-east-2").strip()
_DRAFTS_BUCKET = (os.getenv("S3_BUCKET") or os.getenv("TENANT_BUCKET") or "inksuite-data").strip()
_tenant_base = (os.getenv("TENANT_PREFIX") or "tenants/marble-press").strip().rstrip("/")
_DRAFTS_PREFIX = (os.getenv("DRAFTS_S3_PREFIX") or f"{_tenant_base}/data/TempDraftContracts").strip().rstrip("/") + "/"


def _drafts_s3_endpoint() -> str:
    return f"https://s3.{_DRAFTS_AWS_REGION}.amazonaws.com"


def _drafts_s3_client():
    return boto3.client(
        "s3",
        region_name=_DRAFTS_AWS_REGION,
        endpoint_url=_drafts_s3_endpoint(),
        config=Config(
            signature_version="s3v4",
            retries={"max_attempts": 5, "mode": "standard"},
            s3={"addressing_style": "virtual"},
        ),
    )


# ----------------------- PATHS -----------------------
BASE_DATA_DIR = Path.home() / "Documents" / "marble_app" / "data"
TEMPLATES_DIR = BASE_DATA_DIR / "Templates"
DRAFTS_DIR = BASE_DATA_DIR / "TempDraftContracts"
INDEX_PATH = DRAFTS_DIR / "index.json"

if _S3_AVAILABLE:
    DRAFTS_S3_PREFIX = tenant_data_prefix("data", "TempDraftContracts").rstrip("/") + "/"
    DRAFT_INDEX_KEY = DRAFTS_S3_PREFIX + "index.json"
    TEMPLATES_S3_PREFIX = tenant_data_prefix("data", "Templates").rstrip("/") + "/"
    TEMPLATES_INDEX_KEY = TEMPLATES_S3_PREFIX + "templates_index.json"
else:
    DRAFTS_S3_PREFIX = ""
    DRAFT_INDEX_KEY = ""
    TEMPLATES_S3_PREFIX = ""
    TEMPLATES_INDEX_KEY = ""

TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
if not INDEX_PATH.exists():
    INDEX_PATH.write_text("[]", encoding="utf-8")

TOKEN_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

SUBRIGHT_TOKENS = [
    "Sub_BookClub",
    "Sub_FirstSerial_Illustrated",
    "Sub_FirstSerial_Text",
    "Sub_SecondSerial",
    "Sub_Audio_Physical",
    "Sub_Audio_Digital",
    "Sub_UK",
    "Sub_Canada",
    "Sub_Export",
    "Sub_ForeignTranslation",
    "Sub_TV_Movie",
    "Sub_Animation_Digital",
    "Sub_MassMerch",
]

_last_draft_list_error: Optional[str] = None


# ----------------------- helpers -----------------------

def _as_bool(v: Any) -> bool:
    if isinstance(v, bool):
        return v
    if v is None:
        return False
    if isinstance(v, (int, float)):
        return v != 0
    if isinstance(v, str):
        s = v.strip().lower()
        if s in {"true", "1", "yes", "y", "on"}:
            return True
        if s in {"false", "0", "no", "n", "off", ""}:
            return False
    return False

def _normalize_token_name(token: str) -> str:
    return re.sub(r"\s+", " ", str(token or "").strip())


def _dig(obj, dotted: str):
    cur = obj
    for part in dotted.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return None
    return cur


def _first_non_empty(*vals):
    for v in vals:
        if v is None:
            continue
        if isinstance(v, str) and v.strip() == "":
            continue
        return v
    return ""


def _trim(v: Any) -> str:
    return "" if v is None else str(v).strip()


def _decimal_or_zero(v: Any) -> Decimal:
    try:
        if v in (None, ""):
            return Decimal("0")
        return Decimal(str(v))
    except Exception:
        return Decimal("0")


def _money(v: Any) -> str:
    try:
        n = _decimal_or_zero(v).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        return f"${n:,.2f}"
    except Exception:
        return "$0.00"


def _fmt_percent(v: Any) -> str:
    try:
        d = _decimal_or_zero(v)
        s = format(d.normalize(), "f")
        if "." in s:
            s = s.rstrip("0").rstrip(".")
        return s or "0"
    except Exception:
        return _trim(v) or "0"


def _append_text_to_paragraph(p, text: str, color: RGBColor | None = None) -> None:
    text = "" if text is None else str(text)
    parts = text.split("\n")

    last_run = None
    for i, part in enumerate(parts):
        if i == 0:
            run = p.add_run(part)
            if color is not None:
                run.font.color.rgb = color
            last_run = run
            continue

        if last_run is None:
            last_run = p.add_run("")
            if color is not None:
                last_run.font.color.rgb = color

        last_run.add_break()
        run = p.add_run(part)
        if color is not None:
            run.font.color.rgb = color
        last_run = run


def _load_deal_memo_db_fields(uid: str) -> dict[str, Any]:
    if not uid:
        return {}

    try:
        with db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT
                        d.*,
                        COALESCE(agent.display_name, '') AS agent_name,
                        COALESCE(agent.email, '') AS agent_email,
                        COALESCE(agent.website, '') AS agent_website,
                        COALESCE(agent.phone_country_code, '') AS agent_phone_country_code,
                        COALESCE(agent.phone_number, '') AS agent_phone_number,

                        COALESCE(agency.display_name, '') AS agency_name,
                        COALESCE(agency.email, '') AS agency_email,
                        COALESCE(agency.website, '') AS agency_website,
                        COALESCE(agency.phone_country_code, '') AS agency_phone_country_code,
                        COALESCE(agency.phone_number, '') AS agency_phone_number,

                        COALESCE(addr.street, '') AS agency_street,
                        COALESCE(addr.city, '') AS agency_city,
                        COALESCE(addr.state, '') AS agency_state,
                        COALESCE(addr.zip, '') AS agency_zip,
                        COALESCE(addr.country, '') AS agency_country,

                        COALESCE(aprof.agency_clause, '') AS agency_clause
                    FROM deal_memo_drafts d
                    LEFT JOIN parties agent
                        ON agent.id = d.agent_party_id
                    LEFT JOIN parties agency
                        ON agency.id = d.agency_party_id
                    LEFT JOIN agency_profiles aprof
                        ON aprof.agency_party_id = d.agency_party_id
                    LEFT JOIN party_addresses addr
                        ON addr.party_id = d.agency_party_id
                    WHERE d.uid = %s
                    LIMIT 1
                    """,
                    (uid,),
                )
                row = cur.fetchone()
                if not row:
                    return {}

                cols = [desc[0] for desc in cur.description]
                out = dict(zip(cols, row))

                out["author_address"] = {
                    "street": out.get("author_street") or "",
                    "city": out.get("author_city") or "",
                    "state": out.get("author_state") or "",
                    "zip": out.get("author_zip") or "",
                    "country": out.get("author_country") or "",
                }
                out["illustrator_address"] = {
                    "street": out.get("illustrator_street") or "",
                    "city": out.get("illustrator_city") or "",
                    "state": out.get("illustrator_state") or "",
                    "zip": out.get("illustrator_zip") or "",
                    "country": out.get("illustrator_country") or "",
                }

                out["authorAgent"] = {
                    "name": out.get("agent_name") or "",
                    "agency": out.get("agency_name") or "",
                    "email": out.get("agent_email") or "",
                    "website": out.get("agency_website") or "",
                    "address": {
                        "street": out.get("agency_street") or "",
                        "city": out.get("agency_city") or "",
                        "state": out.get("agency_state") or "",
                        "zip": out.get("agency_zip") or "",
                        "country": out.get("agency_country") or "",
                    },
                }
                out["author_agent"] = out["authorAgent"]

                return out
    except Exception:
        return {}


def _get_value_from_memo(memo: dict, dotted: str) -> str:
    val = _dig(memo, dotted)
    if val not in (None, ""):
        return "" if val is None else str(val)

    aliases = {
        "contributor_role": ["contributorRole"],

        "author_email": ["authorEmail"],
        "author_phone_number": ["authorPhoneNumber"],
        "author_phone_country_code": ["authorPhoneCountryCode"],

        "illustrator_name": ["illustrator.name", "illustratorName"],
        "illustrator_email": ["illustrator.email", "illustratorEmail"],
        "illustrator_phone_number": ["illustrator.phoneNumber", "illustratorPhoneNumber"],

        "author_address.street": ["authorAddress.street", "author_street"],
        "author_address.city": ["authorAddress.city", "author_city"],
        "author_address.state": ["authorAddress.state", "author_state"],
        "author_address.zip": ["authorAddress.zip", "authorAddress.postalCode", "author_zip"],
        "author_address.country": ["authorAddress.country", "author_country"],

        "illustrator_address.street": ["illustratorAddress.street", "illustrator.address.street", "illustrator_street"],
        "illustrator_address.city": ["illustratorAddress.city", "illustrator.address.city", "illustrator_city"],
        "illustrator_address.state": ["illustratorAddress.state", "illustrator.address.state", "illustrator_state"],
        "illustrator_address.zip": [
            "illustratorAddress.zip",
            "illustrator.address.zip",
            "illustratorAddress.postalCode",
            "illustrator_zip",
        ],
        "illustrator_address.country": ["illustratorAddress.country", "illustrator.address.country", "illustrator_country"],

        "agent_name": [
            "agentName",
            "authors_agent_name",
            "authorsAgentName",
            "authorAgent.name",
            "author_agent.name",
        ],
        "agency_name": [
            "agencyName",
            "authorAgent.agency",
            "author_agent.agency",
            "authorAgent.name",
            "author_agent.name",
        ],
        "agent_email": [
            "agentEmail",
            "authors_agent_email",
            "authorsAgentEmail",
            "authorAgent.email",
            "author_agent.email",
        ],
        "agency_email": ["agencyEmail"],
        "agency_website": [
            "agencyWebsite",
            "authorAgent.website",
            "author_agent.website",
        ],
        "agency_street": [
            "agencyStreet",
            "authorAgent.address.street",
            "author_agent.address.street",
        ],
        "agency_city": [
            "agencyCity",
            "authorAgent.address.city",
            "author_agent.address.city",
        ],
        "agency_state": [
            "agencyState",
            "authorAgent.address.state",
            "author_agent.address.state",
        ],
        "agency_zip": [
            "agencyZip",
            "authorAgent.address.zip",
            "author_agent.address.zip",
            "authorAgent.address.postalCode",
        ],
        "agency_country": [
            "agencyCountry",
            "authorAgent.address.country",
            "author_agent.address.country",
        ],
        "agency_clause": ["agencyClause"],

        "effective_date": ["effectiveDate"],
        "projected_publication_date": ["projectedPublicationDate"],
        "projected_retail_price": ["projectedRetailPrice"],
        "territories_rights": ["territoriesRights"],
        "short_description": ["shortDescription"],
        "option_clause": ["optionClause"],

        "author_advance": ["authorAdvance", "advance", "totalAdvance"],
        "illustrator_advance": ["illustratorAdvance"],

        "comp_copies_contributor": ["compCopiesContributor"],
        "comp_copies_agent": ["compCopiesAgent"],

        "delivery_mode": ["deliveryMode"],
        "delivery_clause": ["deliveryClause"],
        "delivery_date": ["deliveryDate"],

        "advance_schedule": ["advanceSchedule"],
    }

    candidates = [dotted] + aliases.get(dotted, [])
    vals = [_dig(memo, c) for c in candidates]
    val = _first_non_empty(*vals)
    return "" if val is None else str(val)


def _default_mapping(role: str) -> dict[str, str]:
    is_illustrator = (role == "illustrator")

    return {
        "Author_Name": "illustrator_name" if is_illustrator else "author",
        "Author_Email": "illustrator_email" if is_illustrator else "author_email",
        "Author_Phone": "illustrator_phone_number" if is_illustrator else "author_phone_number",

        "AUTHOR_NAME": "illustrator_name" if is_illustrator else "author",
        "AUTHOR_EMAIL": "illustrator_email" if is_illustrator else "author_email",
        "AUTHOR_PHONE": "illustrator_phone_number" if is_illustrator else "author_phone_number",

        "Author_Street_Address": "illustrator_address.street" if is_illustrator else "author_address.street",
        "Author_City": "illustrator_address.city" if is_illustrator else "author_address.city",
        "Author_State": "illustrator_address.state" if is_illustrator else "author_address.state",
        "Author_Zip": "illustrator_address.zip" if is_illustrator else "author_address.zip",

        "AUTHOR_STREET_ADDRESS": "illustrator_address.street" if is_illustrator else "author_address.street",
        "AUTHOR_CITY": "illustrator_address.city" if is_illustrator else "author_address.city",
        "AUTHOR_STATE": "illustrator_address.state" if is_illustrator else "author_address.state",
        "AUTHOR_ZIP": "illustrator_address.zip" if is_illustrator else "author_address.zip",

        "Date": "effective_date",
        "Book_Title": "title",
        "Book_Description": "short_description",
        "Option_Clause": "option_clause",
        "Projected_Publication": "projected_publication_date",
        "Territory": "territories_rights",
        "Right Limitation": "territories_rights",

        "DATE": "effective_date",
        "BOOK_TITLE": "title",
        "BOOK_DESCRIPTION": "short_description",
        "OPTION_CLAUSE": "option_clause",
        "PROJECTED_PUBLICATION": "projected_publication_date",
        "TERRITORY": "territories_rights",

        "Comp_Copies_Contributor": "comp_copies_contributor",
        "Comp_Copies_Agent": "comp_copies_agent",
        "AUTHOR_COPIES": "comp_copies_contributor",
        "AGENCY_COPIES": "comp_copies_agent",

        "Total_Advance": "illustrator_advance" if is_illustrator else "author_advance",
        "TOTAL_ADVANCE": "illustrator_advance" if is_illustrator else "author_advance",
        "total advence": "illustrator_advance" if is_illustrator else "author_advance",
        "TOTAL_ADVENCE": "illustrator_advance" if is_illustrator else "author_advance",

        "Advance_Installments": "__ADVANCE_INSTALLMENTS__",
        "Advance_Installments_Block": "__ADVANCE_INSTALLMENTS_BLOCK__",
        "Advance_Installments_Sentence": "__ADVANCE_INSTALLMENTS_SENTENCE__",

        "Agent Name": "agent_name",
        "Agency Name": "agency_name",
        "Agent Email": "agent_email",
        "Agency Email": "agency_email",
        "Agency Website": "agency_website",
        "Agency Address": "agency_street",
        "Agency Street": "agency_street",
        "Agency Street Address": "agency_street",
        "Agency City": "agency_city",
        "Agency State": "agency_state",
        "Agency Zip": "agency_zip",
        "Agency Country": "agency_country",
        "Agency_Clause": "agency_clause",

        "AGENT_NAME": "agent_name",
        "AGENCY_NAME": "agency_name",
        "AGENT_EMAIL": "agent_email",
        "AGENCY_EMAIL": "agency_email",
        "AGENCY_WEBSITE": "agency_website",
        "AGENCY_ADDRESS": "agency_street",
        "AGENCY_STREET": "agency_street",
        "AGENCY_STREET_ADDRESS": "agency_street",
        "AGENCY_CITY": "agency_city",
        "AGENCY_STATE": "agency_state",
        "AGENCY_ZIP": "agency_zip",
        "AGENCY_COUNTRY": "agency_country",
        "AGENCY_CLAUSE": "agency_clause",

        "Projected_Retail_Price": "projected_retail_price",
        "Delivery_Mode": "delivery_mode",
        "Delivery_Clause": "delivery_clause",
        "Delivery_Date": "delivery_date",

        "PROJECTED_RETAIL_PRICE": "projected_retail_price",
        "DELIVERY_MODE": "delivery_mode",
        "DELIVERY_CLAUSE": "delivery_clause",
        "DELIVERY_DATE": "delivery_date",

        "Manuscript_Delivery_Block": "__MANUSCRIPT_DELIVERY_BLOCK__",
        "MANUSCRIPT_DELIVERY_BLOCK": "__MANUSCRIPT_DELIVERY_BLOCK__",
    }


def _load_index() -> List[dict]:
    if _S3_AVAILABLE and DRAFT_INDEX_KEY:
        try:
            raw = get_bytes(DRAFT_INDEX_KEY)
            data = json.loads(raw.decode("utf-8"))
            if isinstance(data, list):
                return data
            if isinstance(data, dict) and isinstance(data.get("items"), list):
                return data["items"]
        except Exception:
            pass
    try:
        return json.loads(INDEX_PATH.read_text(encoding="utf-8"))
    except Exception:
        return []


def _save_index(items: List[dict]) -> None:
    if _S3_AVAILABLE and DRAFT_INDEX_KEY:
        try:
            put_bytes(
                DRAFT_INDEX_KEY,
                json.dumps(items, ensure_ascii=False, indent=2).encode("utf-8"),
                content_type="application/json",
            )
            return
        except Exception:
            pass
    INDEX_PATH.write_text(json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8")


def _list_draft_keys_direct() -> List[str]:
    global _last_draft_list_error
    _last_draft_list_error = None
    keys: List[str] = []
    if not _DRAFTS_BUCKET or not _DRAFTS_PREFIX:
        _last_draft_list_error = "DRAFTS_BUCKET or DRAFTS_PREFIX empty"
        return keys
    try:
        client = _drafts_s3_client()
        token: Optional[str] = None
        while True:
            kwargs: Dict[str, Any] = {"Bucket": _DRAFTS_BUCKET, "Prefix": _DRAFTS_PREFIX, "MaxKeys": 1000}
            if token:
                kwargs["ContinuationToken"] = token
            resp = client.list_objects_v2(**kwargs)
            for obj in resp.get("Contents") or []:
                k = (obj.get("Key") or "").strip()
                if k and not k.endswith("/") and k.lower().endswith(".docx"):
                    keys.append(k)
            if not resp.get("IsTruncated"):
                break
            token = resp.get("NextContinuationToken") or None
            if not token:
                break
    except Exception as e:
        _last_draft_list_error = f"{type(e).__name__}: {e}"
    return keys


def _synthesize_drafts_from_s3() -> List[dict]:
    out: List[dict] = []
    keys = _list_draft_keys_direct()

    if not keys and _S3_AVAILABLE and DRAFTS_S3_PREFIX:
        try:
            keys = list_keys(DRAFTS_S3_PREFIX)
            keys = [k for k in keys if k and k.lower().endswith(".docx")]
        except Exception:
            pass

    for key in keys:
        basename = key.split("/")[-1] if "/" in key else key
        if not basename or basename == "index.json" or not basename.lower().endswith(".docx"):
            continue
        stem = basename[:-5]
        out.append(
            {
                "id": stem,
                "uid": stem,
                "title": stem.replace("_", " ").strip() or stem,
                "filename": basename,
                "s3_key": key,
                "path": "",
                "templateId": "",
                "createdAt": datetime.utcnow().isoformat() + "Z",
            }
        )
    return sorted(out, key=lambda x: x.get("createdAt", ""), reverse=True)


def _find_template_path(template_id: str) -> Path:
    if template_id.lower().endswith(".docx"):
        p = TEMPLATES_DIR / template_id
        if p.exists():
            return p
        raise HTTPException(status_code=404, detail=f"Template file not found: {p}")

    matches = sorted(TEMPLATES_DIR.glob(f"{template_id}_*.docx"))
    if matches:
        return matches[0]

    fallback = TEMPLATES_DIR / f"{template_id}.docx"
    if fallback.exists():
        return fallback

    raise HTTPException(
        status_code=404,
        detail=(
            "Template not found.\n"
            f"Looked for:\n - {TEMPLATES_DIR}\\{template_id}_*.docx\n - {TEMPLATES_DIR}\\{template_id}.docx"
        ),
    )


def _load_templates_index() -> list[dict[str, Any]]:
    if _S3_AVAILABLE and TEMPLATES_INDEX_KEY:
        try:
            raw = get_bytes(TEMPLATES_INDEX_KEY)
            data = json.loads(raw.decode("utf-8"))
            if isinstance(data, list):
                return data
            if isinstance(data, dict) and isinstance(data.get("items"), list):
                return data["items"]
        except Exception:
            pass

    local_index = TEMPLATES_DIR / "templates_index.json"
    try:
        data = json.loads(local_index.read_text(encoding="utf-8"))
        if isinstance(data, list):
            return data
        if isinstance(data, dict) and isinstance(data.get("items"), list):
            return data["items"]
    except Exception:
        pass

    return []


def _resolve_template_s3_key(template_id: str) -> str | None:
    items = _load_templates_index()

    for it in items:
        if str(it.get("id") or "") == str(template_id):
            for key_name in ("s3_key", "key", "path", "object_key"):
                key = str(it.get(key_name) or "").strip()
                if key and key.lower().endswith(".docx"):
                    return key

            filename = str(it.get("filename") or "").strip()
            if filename and TEMPLATES_S3_PREFIX:
                return TEMPLATES_S3_PREFIX + filename

    if _S3_AVAILABLE and TEMPLATES_S3_PREFIX:
        try:
            keys = list_keys(TEMPLATES_S3_PREFIX)
            for key in keys:
                base = key.split("/")[-1]
                if base == template_id or base == f"{template_id}.docx":
                    return key
            for key in keys:
                base = key.split("/")[-1]
                if base.startswith(f"{template_id}_") and base.lower().endswith(".docx"):
                    return key
        except Exception:
            pass

    return None


def _load_template_bytes_and_name(template_id: str) -> tuple[bytes, str]:
    s3_key = _resolve_template_s3_key(template_id)
    if s3_key and _S3_AVAILABLE:
        try:
            data = get_bytes(s3_key)
            filename = s3_key.split("/")[-1]
            return data, filename
        except Exception:
            pass

    path = _find_template_path(template_id)
    return path.read_bytes(), path.name


def _populate_royalty_tokens(memo: dict, values: dict[str, str]) -> bool:
    def _fmt(v) -> str:
        return "" if v is None else str(v)

    royalties = (memo.get("royalties") or {}).get("author") or {}
    first_rights = royalties.get("first_rights") or []

    def _normalize_fmt_name(fmt: str) -> str:
        return fmt.lower().replace("-", "").replace(" ", "")

    def _find_format(fmt: str):
        target = _normalize_fmt_name(fmt)
        for r in first_rights:
            name = _normalize_fmt_name(str(r.get("format") or ""))
            if name == target:
                return r
        return None

    has_boardbook = False

    def _fill_format(prefix: str, fmt_key: str) -> bool:
        nonlocal has_boardbook
        fr = _find_format(fmt_key)
        if not fr:
            return False

        tiers = fr.get("tiers") or []
        if not tiers:
            return False

        t1 = tiers[0]
        values[f"{prefix}_1"] = _fmt(t1.get("rate_percent"))

        copy_limit = None
        for cond in t1.get("conditions") or []:
            if cond.get("kind") == "units" and cond.get("value") is not None:
                copy_limit = cond["value"]
                break
        if copy_limit is not None:
            values[f"{prefix}_Copy_Limit_1"] = _fmt(copy_limit)

        if len(tiers) > 1:
            values[f"{prefix}_2"] = _fmt(tiers[1].get("rate_percent"))
        if len(tiers) > 2:
            values[f"{prefix}_3"] = _fmt(tiers[2].get("rate_percent"))
        if len(tiers) > 3:
            values[f"{prefix}_4"] = _fmt(tiers[3].get("rate_percent"))

        if prefix.lower().startswith("boardbook"):
            has_boardbook = True

        return True

    _fill_format("Hardcover", "Hardcover")
    _fill_format("Paperback", "Paperback")
    _fill_format("Boardbook", "Board Book")

    ebook = _find_format("E-book") or _find_format("Ebook")
    if ebook:
        values["Ebook"] = _fmt(ebook.get("flat_rate_percent"))

    return has_boardbook


def _populate_subrights(memo: dict, values: dict[str, str]) -> None:
    royalties = (memo.get("royalties") or {}).get("author") or {}
    subrights = royalties.get("subrights") or []
    if not subrights:
        values["Sub_Canada"] = "2/3"
        values["Sub_Export"] = "2/3"
        return

    def norm(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", (s or "").lower()).strip()

    def find_sub(fragment: str):
        frag = fragment.lower()
        for item in subrights:
            if frag in norm(item.get("name", "")):
                return item
        return None

    def set_percent(token: str, item: dict | None, variant_key: str | None = None):
        if not item:
            return
        if variant_key:
            v = item.get("variants") or {}
            if variant_key in v and v[variant_key] is not None:
                values[token] = str(v[variant_key])
        else:
            p = item.get("percent")
            if p is not None:
                values[token] = str(p)

    set_percent("Sub_HC_PB_LargeType", find_sub("hardcover paperback"))
    set_percent("Sub_Anthologies", find_sub("anthologies"))
    set_percent("Sub_BookClub", find_sub("book club"))

    first_serial = find_sub("first serial")
    set_percent("Sub_FirstSerial_Text", first_serial, variant_key="text_only")
    set_percent("Sub_FirstSerial_Illustrated", first_serial, variant_key="text_and_art")

    second_serial = find_sub("second serial")
    set_percent("Sub_SecondSerial", second_serial, variant_key="text_only")

    audio = find_sub("audiobook") or find_sub("audiobooks")
    set_percent("Sub_Audio_Physical", audio, variant_key="physical")
    set_percent("Sub_Audio_Digital", audio, variant_key="digital")

    set_percent("Sub_UK", find_sub("uk"))
    values["Sub_Canada"] = "2/3"
    values["Sub_Export"] = "2/3"
    set_percent("Sub_ForeignTranslation", find_sub("foreign translation"))


def _find_draft(item_id: str) -> dict | None:
    for it in _load_index():
        if str(it.get("id")) == str(item_id):
            return it
    for it in _synthesize_drafts_from_s3():
        if str(it.get("id")) == str(item_id):
            return it
    return None


def _load_advance_installments_db(uid: str) -> list[dict[str, Any]]:
    if not uid:
        return []

    try:
        with db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT
                        ai.installment_order,
                        ai.amount_type,
                        ai.value,
                        ai.trigger
                    FROM deal_memo_advance_installments ai
                    JOIN deal_memo_drafts d
                        ON d.id = ai.deal_memo_draft_id
                    WHERE d.uid = %s
                    ORDER BY ai.installment_order, ai.id
                    """,
                    (uid,),
                )
                rows = cur.fetchall() or []
                cols = [desc[0] for desc in cur.description]
                return [dict(zip(cols, row)) for row in rows]
    except Exception:
        return []


def _build_advance_installments_block(total_advance: Any, rows: list[dict[str, Any]]) -> str:
    total = _decimal_or_zero(total_advance)
    lines: list[str] = []

    for idx, row in enumerate(rows or [], start=1):
        amount_type = _trim(row.get("amount_type")).lower()
        value = row.get("value")
        trigger = _trim(row.get("trigger"))

        value_dec = _decimal_or_zero(value)

        if amount_type == "percent":
            dollars = (total * value_dec / Decimal("100")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            amount_text = f"{_fmt_percent(value_dec)}% of the total advance ({_money(dollars)})"
        else:
            dollars = value_dec.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            amount_text = _money(dollars)

        line = f"{idx}. {amount_text}"
        if trigger:
            line += f", {trigger}"
        line += ";"
        lines.append(line)

    if not lines:
        return ""

    lines[-1] = lines[-1].rstrip(";") + "."
    return "\n".join(lines)


def _build_advance_installments_sentence(total_advance: Any, rows: list[dict[str, Any]]) -> str:
    total = _decimal_or_zero(total_advance)
    parts: list[str] = []

    for row in rows or []:
        amount_type = _trim(row.get("amount_type")).lower()
        value = row.get("value")
        trigger = _trim(row.get("trigger"))

        value_dec = _decimal_or_zero(value)

        if amount_type == "percent":
            dollars = (total * value_dec / Decimal("100")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            amount_text = f"{_fmt_percent(value_dec)}% of the total advance ({_money(dollars)})"
        else:
            dollars = value_dec.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            amount_text = _money(dollars)

        text = amount_text
        if trigger:
            text += f" {trigger}"
        parts.append(text.strip())

    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2:
        return f"{parts[0]}; and {parts[1]}"
    return "; ".join(parts[:-1]) + f"; and {parts[-1]}"


def _build_delivery_follow_on_package() -> str:
    return (
        "b) The Work shall be deemed accepted by Publisher only if Publisher has expressly confirmed its "
        "acceptance, in writing, to Author, and expressions of encouragement or payment of money shall not "
        "be deemed to constitute acceptance. Publisher shall advise Author in writing within forty-five (45) "
        "days of its receipt of the complete Work whether or not the Work is acceptable in length, content, "
        "and form to Publisher. If Publisher fails to so notify Author within such forty-five (45) day period, "
        "Author shall give Publisher written notification of Publisher’s failure to do so, whereupon Publisher "
        "shall have an additional thirty (30) days from receipt of such written notification in which to inform "
        "Author whether or not the Work is editorially acceptable. If Publisher still fails to so notify Author "
        "after such additional thirty (30) day period, the Work shall be deemed automatically accepted.\n\n"
        "c) If the Work as delivered is not editorially satisfactory to Publisher, Publisher shall provide Author "
        "with reasonably detailed written suggestions for revisions and request Author to work cooperatively with "
        "Publisher to make the Work satisfactory to Publisher within a time period to be fixed by Publisher, which "
        "time period shall be reasonably related to the requested revisions and shall be at least thirty (30) days "
        "in length in any event, and Author shall use Author’s best efforts to do so.\n\n"
        "d) If Author fails to revise the Work so as to render it editorially satisfactory to Publisher in "
        "accordance with subparagraph (c), then Publisher shall have the right to terminate this Agreement, in "
        "which event Author may offer the rights to the Work elsewhere. If a third-party publisher thereafter "
        "wishes to acquire publication rights to the Work, Author shall submit to Publisher written agreement for "
        "repayment to Publisher of the actual amount advanced by Publisher to Author under this Agreement out of "
        "the first monies to be paid to Author for the Work by such third party within eighteen (18) months from "
        "Publisher’s notice after termination. Publisher shall terminate and relinquish its rights with respect to "
        "the Work conditioned upon Author entering into such agreement with the third party, and in any event, "
        "whether or not Author enters into agreement with a third party for the Work, Author agrees to repay the "
        "Advance for the Work to Publisher within such eighteen (18) months from notification that the Work is not "
        "editorially acceptable.\n\n"
        "e) If Author fails to comply with the requirements of Paragraph 11, then Publisher shall have the option "
        "itself to obtain the permissions and materials referred to in Paragraph 11 and charge the commercially "
        "reasonable and documented costs thereof against any sums payable to Author.\n\n"
        "f) In no event shall Publisher engage another author to complete the Work."
    )


def _build_manuscript_delivery_block(memo: dict) -> str:
    role = str(memo.get("contributorRole") or memo.get("contributor_role") or "author").lower()
    mode = memo.get("deliveryMode") or memo.get("delivery_mode") or "author_signing"

    if role == "illustrator":
        return (memo.get("deliveryClause") or memo.get("delivery_clause") or "").strip()

    if mode == "author_done":
        return "1. Manuscript has been delivered, editorially satisfactory, and accepted."

    if mode == "author_signing":
        return "1. Manuscript shall be delivered to Publisher in an editorially satisfactory format upon signing."

    if mode == "author_by_date":
        lead = (
            f"a) The final manuscript text shall be delivered to Publisher not later than "
            f"{memo.get('deliveryDate') or memo.get('delivery_date') or '{{date}}'}."
        )
        return lead + "\n\n" + _build_delivery_follow_on_package()

    if mode == "author_series":
        rows = memo.get("seriesDeliveries") or memo.get("series_deliveries") or []

        parts = []
        for idx, row in enumerate(rows, start=1):
            book = row.get("book") or f"Book {idx}"
            date = row.get("byDate") or row.get("by_date") or "{{date}}"
            parts.append(f"{chr(96 + idx)}) The final manuscript text of {book} shall be delivered to Publisher not later than {date}.")

        lead = "\n".join(parts) if parts else (
            "a) The final manuscript texts shall be delivered to Publisher according to a mutually agreed schedule."
        )

        return lead + "\n\n" + _build_delivery_follow_on_package()

    return "1. Manuscript shall be delivered to Publisher in an editorially satisfactory format upon signing."


def _wopi_settings():
    if get_settings is None:
        raise HTTPException(status_code=500, detail="WOPI not configured")
    return get_settings()


def _extract_wopi_token(request: Request) -> str:
    token = request.query_params.get("access_token")
    if token:
        return token
    auth = request.headers.get("Authorization")
    if auth and auth.startswith("Bearer "):
        return auth[7:].strip()
    return request.headers.get("X-WOPI-AccessToken") or ""


def _put_draft_contract_file(item_id: str, content: bytes) -> None:
    it = _find_draft(item_id)
    if not it:
        raise HTTPException(status_code=404, detail="Draft not found")
    s3_key = it.get("s3_key")
    if s3_key and _DRAFTS_BUCKET:
        client = _drafts_s3_client()
        client.put_object(
            Bucket=_DRAFTS_BUCKET,
            Key=s3_key,
            Body=content,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        return
    p = Path(it.get("path", ""))
    if not p.suffix and not str(p).endswith(".docx"):
        p = p.with_suffix(".docx")
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_bytes(content)


def _wopi_file_info(file_id: str) -> tuple[str, int, str]:
    it = _find_draft(file_id)
    if not it:
        raise HTTPException(status_code=404, detail="File not found")
    base_name = it.get("filename") or (it.get("s3_key") or "").split("/")[-1] or Path(it.get("path", "")).name or "draft.docx"
    size = 0
    s3_key = it.get("s3_key")
    if s3_key and _DRAFTS_BUCKET:
        try:
            client = _drafts_s3_client()
            head = client.head_object(Bucket=_DRAFTS_BUCKET, Key=s3_key)
            size = int(head.get("ContentLength", 0))
        except Exception:
            pass
    if size == 0:
        p = Path(it.get("path", ""))
        if p.exists():
            size = p.stat().st_size
        else:
            try:
                size = len(__wopi_get_file_bytes(file_id))
            except Exception:
                pass
    version = str(it.get("updated_at") or it.get("createdAt") or id(it))
    return base_name, size, version


def __wopi_get_file_bytes(file_id: str) -> bytes:
    it = _find_draft(file_id)
    if not it:
        raise HTTPException(status_code=404, detail="File not found")
    s3_key = it.get("s3_key")
    if s3_key and _DRAFTS_BUCKET:
        client = _drafts_s3_client()
        resp = client.get_object(Bucket=_DRAFTS_BUCKET, Key=s3_key)
        return resp["Body"].read()
    p = Path(it.get("path", ""))
    if not p.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return p.read_bytes()


# ------------------------ WOPI ------------------------

@wopi_router.get("/files/{file_id}")
def wopi_check_file_info(file_id: str, request: Request):
    token = _extract_wopi_token(request)
    settings = _wopi_settings()
    payload = verify_wopi_token(token, settings.wopi_access_token_secret)
    if str(payload.get("file_id")) != str(file_id):
        raise HTTPException(status_code=401, detail="Invalid WOPI token")
    base_name, size, version = _wopi_file_info(file_id)
    user_id = (payload.get("user_id") or "inksuite-user").strip()
    user_id = "".join(c if c.isalnum() else "_" for c in user_id)[:64] or "inksuite_user"
    return JSONResponse(
        {
            "BaseFileName": base_name,
            "Size": size,
            "OwnerId": "inksuite",
            "UserId": user_id,
            "UserFriendlyName": payload.get("user_id") or "InkSuite User",
            "Version": version,
            "SupportsUpdate": True,
            "SupportsLocks": True,
            "SupportsGetLock": True,
            "SupportsPutFile": True,
            "SupportsRename": False,
            "ReadOnly": False,
            "UserCanWrite": True,
        }
    )


@wopi_router.get("/files/{file_id}/contents")
def wopi_get_file(file_id: str, request: Request):
    token = _extract_wopi_token(request)
    settings = _wopi_settings()
    payload = verify_wopi_token(token, settings.wopi_access_token_secret)
    if str(payload.get("file_id")) != str(file_id):
        raise HTTPException(status_code=401, detail="Invalid WOPI token")
    data = __wopi_get_file_bytes(file_id)
    it = _find_draft(file_id)
    base_name = it.get("filename") or "draft.docx"
    return FastAPIResponse(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'inline; filename="{base_name}"'},
    )


@wopi_router.post("/files/{file_id}/contents")
async def wopi_put_file(file_id: str, request: Request):
    token = _extract_wopi_token(request)
    settings = _wopi_settings()
    payload = verify_wopi_token(token, settings.wopi_access_token_secret)
    if str(payload.get("file_id")) != str(file_id):
        raise HTTPException(status_code=401, detail="Invalid WOPI token")
    body = await request.body()
    _put_draft_contract_file(file_id, body)
    return JSONResponse({"LastModifiedTime": datetime.utcnow().isoformat() + "Z"})


_WOPI_LOCKS: dict[str, dict[str, Any]] = {}
_WOPI_LOCK_TTL_SEC = 30 * 60


def _now_epoch() -> float:
    return time.time()


def _lock_get(file_id: str) -> str:
    ent = _WOPI_LOCKS.get(file_id)
    if not ent:
        return ""
    if ent.get("exp", 0) < _now_epoch():
        _WOPI_LOCKS.pop(file_id, None)
        return ""
    return str(ent.get("lock") or "")


def _lock_set(file_id: str, lock_value: str) -> None:
    _WOPI_LOCKS[file_id] = {"lock": lock_value, "exp": _now_epoch() + _WOPI_LOCK_TTL_SEC}


def _lock_clear(file_id: str) -> None:
    _WOPI_LOCKS.pop(file_id, None)


@wopi_router.post("/files/{file_id}")
async def wopi_files_override(file_id: str, request: Request):
    token = _extract_wopi_token(request)
    settings = _wopi_settings()
    payload = verify_wopi_token(token, settings.wopi_access_token_secret)
    if str(payload.get("file_id")) != str(file_id):
        raise HTTPException(status_code=401, detail="Invalid WOPI token")

    override = (request.headers.get("X-WOPI-Override") or "").upper().strip()
    lock_value = request.headers.get("X-WOPI-Lock") or ""

    if not override:
        raise HTTPException(status_code=400, detail="Missing X-WOPI-Override")

    if override == "GET_LOCK":
        cur = _lock_get(file_id)
        resp = FastAPIResponse(content=b"", media_type="text/plain")
        resp.headers["X-WOPI-Lock"] = cur
        return resp

    if override == "LOCK":
        if not lock_value:
            raise HTTPException(status_code=400, detail="Missing X-WOPI-Lock")
        cur = _lock_get(file_id)
        if cur and cur != lock_value:
            resp = FastAPIResponse(status_code=409, content=b"")
            resp.headers["X-WOPI-Lock"] = cur
            return resp
        _lock_set(file_id, lock_value)
        return FastAPIResponse(content=b"")

    if override == "REFRESH_LOCK":
        if not lock_value:
            raise HTTPException(status_code=400, detail="Missing X-WOPI-Lock")
        cur = _lock_get(file_id)
        if cur and cur != lock_value:
            resp = FastAPIResponse(status_code=409, content=b"")
            resp.headers["X-WOPI-Lock"] = cur
            return resp
        _lock_set(file_id, lock_value)
        return FastAPIResponse(content=b"")

    if override == "UNLOCK":
        cur = _lock_get(file_id)
        if cur and lock_value and cur != lock_value:
            resp = FastAPIResponse(status_code=409, content=b"")
            resp.headers["X-WOPI-Lock"] = cur
            return resp
        _lock_clear(file_id)
        return FastAPIResponse(content=b"")

    if override == "PUT":
        body = await request.body()
        if not body:
            raise HTTPException(status_code=400, detail="Empty body on PUT override")
        _put_draft_contract_file(file_id, body)
        resp = JSONResponse({"LastModifiedTime": datetime.utcnow().isoformat() + "Z"})
        resp.headers["X-WOPI-ItemVersion"] = str(int(_now_epoch()))
        return resp

    return JSONResponse({"detail": f"Unsupported X-WOPI-Override: {override}"}, status_code=501)


# ------------------------ Collabora config ------------------------

@router.get("/draft-contracts/{item_id}/collabora-config")
def collabora_config_for_draft(item_id: str):
    it = _find_draft(item_id)
    if not it:
        raise HTTPException(status_code=404, detail="Draft not found")
    settings = _wopi_settings()
    token = make_wopi_token(
        file_id=item_id,
        user_id="inksuite-user",
        ttl=settings.wopi_token_ttl_sec,
        secret=settings.wopi_access_token_secret,
    )
    wopi_base = (os.getenv("WOPI_PUBLIC_BASE") or os.getenv("PUBLIC_BASE_URL") or "").strip().rstrip("/")
    if not wopi_base:
        wopi_base = "http://host.docker.internal:8000"
    wopi_src = f"{wopi_base}/api/wopi/files/{quote(item_id, safe='')}"
    base = settings.collabora_code_url
    if "/loleaflet" in base or "/browser" in base or base.endswith(".html"):
        editor_url = base
    else:
        editor_url = f"{base}/loleaflet/dist/loleaflet.html"
    return JSONResponse(
        {
            "editorUrl": editor_url,
            "wopiSrc": wopi_src,
            "accessToken": token,
            "readOnly": False,
            "debug_wopi_public_base": os.getenv("WOPI_PUBLIC_BASE"),
            "debug_public_base_url": os.getenv("PUBLIC_BASE_URL"),
            "debug_effective_public_base": wopi_base,
        }
    )


@router.get("/drafts/{item_id}/collabora-config")
def collabora_config_for_draft_alias(item_id: str):
    return collabora_config_for_draft(item_id)


# ------------------------ Generate ------------------------

@router.post("/generate")
def generate_contract(req: GenerateRequest):
    incoming_memo = req.dealMemo or {}

    uid = str(incoming_memo.get("uid") or "").strip()
    db_memo = _load_deal_memo_db_fields(uid)

    memo = dict(db_memo or {})
    for k, v in incoming_memo.items():
        if v is None:
            continue
        if isinstance(v, str) and v.strip() == "":
            continue
        memo[k] = v

    role = str(memo.get("contributorRole") or memo.get("contributor_role") or "author").lower()
    is_illustrator = role == "illustrator"

    advance_rows = _load_advance_installments_db(uid)
    total_advance_raw = memo.get("illustrator_advance") if is_illustrator else memo.get("author_advance")
    total_advance_display = _money(total_advance_raw)
    advance_installments_block = _build_advance_installments_block(total_advance_raw, advance_rows)
    advance_installments_sentence = _build_advance_installments_sentence(total_advance_raw, advance_rows)
    agency_clause = _trim(_get_value_from_memo(memo, "agency_clause"))

    raw_option_deleted_front = incoming_memo.get("optionDeleted", None)
    raw_option_deleted_db = memo.get("option_deleted", None)

    if raw_option_deleted_front is not None:
        option_deleted = _as_bool(raw_option_deleted_front)
    else:
        option_deleted = _as_bool(raw_option_deleted_db)

    option_text = _trim(
        incoming_memo.get("optionClause")
        or incoming_memo.get("option_clause")
        or incoming_memo.get("option_supplement")
        or memo.get("optionClause")
        or memo.get("option_clause")
        or memo.get("option_supplement")
    )

    option_value = "__DELETED__" if option_deleted or not option_text else option_text

    try:
        print(sorted(list(memo.keys())))
    except Exception:
        print(memo)

    mapping = _default_mapping("illustrator" if is_illustrator else "author")
    if isinstance(req.mapping, dict):
        mapping.update({k: v for k, v in req.mapping.items() if isinstance(k, str) and isinstance(v, str)})

    
    template_bytes, template_filename = _load_template_bytes_and_name(req.templateId)
    

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(template_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)

        values: dict[str, str] = {}
        for token, key in mapping.items():
            norm_token = _normalize_token_name(token)

            if key == "__ADVANCE_INSTALLMENTS__":
                values[norm_token] = advance_installments_block
            elif key == "__ADVANCE_INSTALLMENTS_BLOCK__":
                values[norm_token] = advance_installments_block
            elif key == "__ADVANCE_INSTALLMENTS_SENTENCE__":
                values[norm_token] = advance_installments_sentence
            elif key == "__MANUSCRIPT_DELIVERY_BLOCK__":
                values[norm_token] = _build_manuscript_delivery_block(memo)
            elif key == "option_clause":
                values[norm_token] = option_value
            else:
                raw_val = _get_value_from_memo(memo, key)

                if key in ("author_advance", "illustrator_advance"):
                    values[norm_token] = total_advance_display if raw_val not in ("", None) else ""
                elif key == "agency_clause":
                    values[norm_token] = agency_clause
                else:
                    values[norm_token] = raw_val

        values["Total_Advance"] = total_advance_display
        values["TOTAL_ADVANCE"] = total_advance_display
        values["total advence"] = total_advance_display
        values["TOTAL_ADVENCE"] = total_advance_display
        values["Total Advance"] = total_advance_display

        values["Advance_Installments"] = advance_installments_block
        values["Advance_Installments_Block"] = advance_installments_block
        values["Advance_Installments_Sentence"] = advance_installments_sentence
        values["Advance Installments"] = advance_installments_block
        values["Advance Installments Block"] = advance_installments_block
        values["Advance Installments Sentence"] = advance_installments_sentence

        values["Agency_Clause"] = agency_clause
        values["AGENCY_CLAUSE"] = agency_clause
        values["Agency Clause"] = agency_clause

        delivery_block = _build_manuscript_delivery_block(memo)
        values["MANUSCRIPT_DELIVERY_BLOCK"] = delivery_block
        values["Manuscript_Delivery_Block"] = delivery_block

        values["Option_Clause"] = option_value
        values["OPTION_CLAUSE"] = option_value
        values["Option Clause"] = option_value

        has_boardbook = _populate_royalty_tokens(memo, values)
        _populate_subrights(memo, values)

        def replace_in_paragraph(p):
            if not p.runs:
                return

            original_text = "".join(r.text for r in p.runs)

            if (not has_boardbook) and ("Boardbook_" in original_text):
                element = p._element
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
                return

            if any(f"{{{{{tok}}}}}" in original_text for tok in SUBRIGHT_TOKENS):
                for tok in SUBRIGHT_TOKENS:
                    if f"{{{{{tok}}}}}" in original_text:
                        val = values.get(tok)
                        if val is None or str(val).strip() == "":
                            element = p._element
                            parent = element.getparent()
                            if parent is not None:
                                parent.remove(element)
                            return

            had_token = False
            START = "\uE000"
            END = "\uE001"

            def repl(m):
                nonlocal had_token
                had_token = True
                tok = _normalize_token_name(m.group(1))
                val = values.get(tok, "") or ""
                return f"{START}{val}{END}"

            new_text = TOKEN_RE.sub(repl, original_text)

            if not had_token:
                return

            p.runs[0].text = ""
            for r in p.runs[1:]:
                r.text = ""

            buf: list[str] = []
            in_token = False

            def flush_normal():
                nonlocal buf
                if buf:
                    _append_text_to_paragraph(p, "".join(buf), color=None)
                    buf = []

            def flush_token():
                nonlocal buf
                if buf:
                    text_val = "".join(buf)

                    if text_val == "__DELETED__":
                        run = p.add_run("[Deleted]")
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                    else:
                        _append_text_to_paragraph(p, text_val, color=RGBColor(0xFF, 0x00, 0x00))

                    buf = []

            for ch in new_text:
                if ch == START:
                    flush_normal()
                    in_token = True
                elif ch == END:
                    flush_token()
                    in_token = False
                else:
                    buf.append(ch)

            if buf:
                if in_token:
                    flush_token()
                else:
                    flush_normal()

        for p in list(doc.paragraphs):
            replace_in_paragraph(p)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in list(cell.paragraphs):
                        replace_in_paragraph(p)

        out_uid = str(memo.get("uid") or uuid.uuid4().hex[:12])
        safe_title = (str(memo.get("title") or memo.get("name") or "Contract").strip() or "Contract").replace(" ", "_")
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_name = f"{safe_title}_{out_uid}_{stamp}.docx"
        out_path = DRAFTS_DIR / out_name
        doc.save(str(out_path))

    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    s3_key: str | None = None
    if _S3_AVAILABLE and _DRAFTS_BUCKET and _DRAFTS_PREFIX:
        try:
            client = _drafts_s3_client()
            data = out_path.read_bytes()
            s3_key = f"{_DRAFTS_PREFIX}{out_name}"
            client.put_object(
                Bucket=_DRAFTS_BUCKET,
                Key=s3_key,
                Body=data,
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception:
            s3_key = None

    items = _load_index()
    item_id = uuid.uuid4().hex[:12]
    item = {
        "id": item_id,
        "uid": out_uid,
        "title": safe_title.replace("_", " "),
        "filename": out_name,
        "path": str(out_path),
        "s3_key": s3_key or "",
        "templateId": req.templateId,
        "createdAt": datetime.now().isoformat(),
    }
    items.insert(0, item)
    _save_index(items)

    try:
        with db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT id FROM tenants WHERE lower(slug) = lower(%s) LIMIT 1", ("marble-press",))
                row = cur.fetchone()
                if not row:
                    cur.execute("SELECT id FROM tenants ORDER BY id LIMIT 1")
                    row = cur.fetchone()
                if row and out_uid:
                    tenant_id = str(row[0]) if hasattr(row, "__getitem__") else str(row["id"])
                    cur.execute(
                        """
                        UPDATE deal_memo_drafts
                        SET generated_contract_filename = %s, generated_at = now(), status = 'generated'
                        WHERE tenant_id = %s AND uid = %s
                        """,
                        (out_name, tenant_id, out_uid),
                    )
            conn.commit()
    except Exception:
        pass

    return {"file": item}


# ------------------------ Draft listing / delete / file ------------------------

def _merged_draft_items() -> List[dict]:
    index_items = _load_index()
    s3_items = _synthesize_drafts_from_s3()
    by_id: Dict[str, dict] = {str(it.get("id", "")): it for it in index_items if it.get("id")}
    for it in s3_items:
        kid = str(it.get("id", ""))
        if kid and kid not in by_id:
            by_id[kid] = it
    out = list(by_id.values())
    out.sort(key=lambda x: (x.get("createdAt") or ""), reverse=True)
    return out


@router.get("/draft-contracts/debug")
def draft_contracts_debug():
    keys = _list_draft_keys_direct()
    merged = _merged_draft_items()
    return {
        "config": {
            "bucket": _DRAFTS_BUCKET,
            "prefix": _DRAFTS_PREFIX,
            "region": _DRAFTS_AWS_REGION,
            "endpoint": _drafts_s3_endpoint(),
        },
        "directListKeysCount": len(keys),
        "directListKeysSample": keys[:10],
        "lastError": _last_draft_list_error,
        "mergedItemsCount": len(merged),
        "mergedItemsSample": [{"id": it.get("id"), "title": it.get("title"), "s3_key": it.get("s3_key")} for it in merged[:5]],
    }


@router.get("/draft-contracts")
def list_draft_contracts():
    return {"items": _merged_draft_items()}


@router.delete("/draft-contracts/{item_id}")
def delete_draft_contract(item_id: str):
    it = _find_draft(item_id)
    if not it:
        raise HTTPException(status_code=404, detail="Draft not found")

    s3_key = (it.get("s3_key") or "").strip()
    if s3_key and _DRAFTS_BUCKET:
        try:
            client = _drafts_s3_client()
            client.delete_object(Bucket=_DRAFTS_BUCKET, Key=s3_key)
        except Exception:
            pass

    raw_path = (it.get("path") or "").strip()
    if raw_path:
        try:
            p = Path(raw_path)
            if p.exists() and p.is_file():
                p.unlink()
        except Exception:
            pass

    items = _load_index()
    items = [x for x in items if str(x.get("id")) != str(item_id)]
    _save_index(items)

    return {"ok": True, "deleted": item_id}


@router.delete("/drafts/{item_id}")
def delete_draft_alias(item_id: str):
    return delete_draft_contract(item_id)


@router.get("/draft-contracts/{item_id}/file")
def get_draft_contract_file(item_id: str):
    it = _find_draft(item_id)
    if not it:
        raise HTTPException(status_code=404, detail="Draft not found")

    s3_key = it.get("s3_key")
    if s3_key and _DRAFTS_BUCKET:
        try:
            client = _drafts_s3_client()
            resp = client.get_object(Bucket=_DRAFTS_BUCKET, Key=s3_key)
            data = resp["Body"].read()
            filename = it.get("filename") or s3_key.split("/")[-1]
            return FastAPIResponse(
                content=data,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'inline; filename="{filename}"'},
            )
        except Exception as e:
            raise HTTPException(status_code=404, detail=f"Draft file not found in S3: {e}")

    p = Path(it.get("path", ""))
    if not p.exists():
        raise HTTPException(status_code=404, detail="Draft file not found on disk")

    return FileResponse(
        str(p),
        filename=p.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/drafts")
def list_drafts_alias():
    return list_draft_contracts()


@router.get("/drafts/{item_id}/file")
def get_draft_file_alias(item_id: str):
    return get_draft_contract_file(item_id)